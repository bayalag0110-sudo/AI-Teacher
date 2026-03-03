import streamlit as st
import streamlit.components.v1 as components
import requests
import pandas as pd
from docx import Document
from io import BytesIO
import PyPDF2
from PIL import Image

# 1. СИСТЕМИЙН ТОХИРГОО
st.set_page_config(page_title="Education", layout="wide", page_icon="🎓")

# --- CUSTOM CSS ---
st.markdown("""
    <style>
    .stApp { background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%); }
    .glass-card {
        background: white !important;
        border-radius: 15px;
        padding: 25px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        margin-bottom: 20px;
        border: 1px solid #e2e8f0;
    }
    .stButton>button {
        width: 100%;
        background: linear-gradient(90deg, #3b82f6, #8b5cf6) !important;
        color: white !important;
        font-weight: 700;
        border-radius: 10px;
    }
    label, p, h1, h2, h3 { color: #1e293b !important; }
    </style>
    """, unsafe_allow_html=True)

# --- ТУСЛАХ ФУНКЦҮҮД ---
def extract_pdf_text(file, start=1, end=3):
    try:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for i in range(start-1, min(end, len(reader.pages))):
            text += reader.pages[i].extract_text()
        return text
    except: return ""

def extract_docx_text(file):
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except: return ""

def create_word_doc(content, title):
    doc = Document()
    doc.add_heading(title, 0)
    # Хүснэгт эсвэл текст хэлбэрээр хадгалах
    doc.add_paragraph(content)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- НЭВТРЭХ ХЭСЭГ ---
if "auth" not in st.session_state: st.session_state.auth = False
if not st.session_state.auth:
    _, col, _ = st.columns([1, 1.2, 1])
    with col:
        st.markdown('<div class="glass-card" style="text-align:center; margin-top:50px;">', unsafe_allow_html=True)
        try: st.image('450633998_2213051369057631_4561852154062620515_n.jpg', width=180)
        except: st.title("🎓 Багшийн туслах")
        u_pwd = st.text_input("🔑 Нууц үг", type="password")
        if st.button("НЭВТРЭХ"):
            if u_pwd == "admin1234": st.session_state.auth = True; st.rerun()
    st.stop()

# --- SIDEBAR МЕНЮ ---
with st.sidebar:
    st.markdown("### 🎓 ЦЭС")
    menu = st.radio("Сонгох:", ["💎 Ээлжит төлөвлөгч", "📝 Тест үүсгэгч", "📝 Даалгавар үүсгэх", "🤖 AI Чатбот", "🌍 Боловсрол"])
    if st.button("🚪 Гарах"): st.session_state.auth = False; st.rerun()

# --- 1. 💎 ЭЭЛЖИТ ТӨЛӨВЛӨГЧ (Боловсролын стандартын дагуу) ---
if menu == "💎 Ээлжит төлөвлөгч":
    st.title("💎 Ээлжит хичээлийн төлөвлөлт")
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        p_file = st.file_uploader("Сурах бичиг (PDF)", type="pdf")
        p_pages = st.slider("Хичээл орох хуудас", 1, 500, (1, 3))
        subj = st.text_input("Хичээл", value="Мэдээллийн технологи")
        grade = st.text_input("Анги", value="9-р анги")
        topic = st.text_input("Хичээлийн сэдэв")
        
        if st.button("🚀 Төлөвлөгөө боловсруулах"):
            if topic:
                with st.spinner("AI стандартын дагуу төлөвлөгөөг боловсруулж байна..."):
                    pdf_context = extract_pdf_text(p_file, p_pages[0], p_pages[1]) if p_file else ""
                    
                    # Стандарт промт
                    prompt = f"""
                    Чи бол Монгол улсын ЕБС-ийн мэргэжлийн багш. Төлөвлөгөөг таны өгсөн 'загвар ээлжит.xlsx' бүтцээр гарга:
                    Хичээл: {subj}, Анги: {grade}, Сэдэв: {topic}
                    PDF-ээс авсан агуулга: {pdf_context[:2000]}
                    
                    1. Суралцахуйн зорилт: Bloom Taxonomy-ийн 6 түвшний дагуу.
                    2. Үйл явцын хүснэгт (Баганууд: Үе шат, Хугацаа, Суралцахуйн үйл ажиллагаа, Багшийн дэмжлэг, Хэрэглэгдэхүүн):
                       - Эхлэл (Зорилго, сэдэлжүүлэг)
                       - Өрнөл (Мэдлэг бүтээх, бүлгийн/ганцаарчилсан ажил)
                       - Төгсгөл (Дүгнэлт, 5 минутын үнэлгээ)
                    3. Даалгавар (3 түвшинд):
                       - Дутуу эзэмшсэн: Нэр томьёо, суурь ойлголт.
                       - Эзэмшиж буй: Жишээ, асуулт.
                       - Бүрэн эзэмшсэн: Бүтээлч бодлого.
                    4. Гэрийн даалгавар ба Ялгаатай сурагчидтай ажиллах аргачлал.
                    """
                    res = requests.post("https://api.groq.com/openai/v1/chat/completions", 
                                       headers={"Authorization": f"Bearer {st.secrets['GROQ_API_KEY']}"},
                                       json={"model": "llama-3.3-70b-versatile", "messages": [{"role": "user", "content": prompt}]})
                    st.session_state.plan_output = res.json()['choices'][0]['message']['content']
        st.markdown('</div>', unsafe_allow_html=True)
        
    with col2:
        if 'plan_output' in st.session_state:
            st.markdown(f'<div class="glass-card">{st.session_state.plan_output}</div>', unsafe_allow_html=True)
            st.download_button("📥 Word файл татах", create_word_doc(st.session_state.plan_output, topic), f"{topic}.docx")

# --- 2. 📝 ТЕСТ ҮҮСГЭГЧ ---
elif menu == "📝 Тест үүсгэгч":
    st.title("📝 Тест боловсруулах")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        t_file = st.file_uploader("Агуулга (PDF)", type="pdf")
        t_pages = st.slider("Тест гаргах хуудас", 1, 500, (1, 2))
        t_count = st.number_input("Асуултын тоо", 5, 20, 10)
        if st.button("🎯 Тест үүсгэх"):
            with st.spinner("AI тест боловсруулж байна..."):
                txt = extract_pdf_text(t_file, t_pages[0], t_pages[1])
                res = requests.post("https://api.groq.com/openai/v1/chat/completions", 
                                   headers={"Authorization": f"Bearer {st.secrets['GROQ_API_KEY']}"},
                                   json={"model": "llama-3.3-70b-versatile", "messages": [{"role": "user", "content": f"Текст: {txt[:3000]}. Энэ агуулгаар {t_count} тест зохиож, хариуг бич."}]})
                st.session_state.test_res = res.json()['choices'][0]['message']['content']
        st.markdown('</div>', unsafe_allow_html=True)
    with col2:
        if 'test_res' in st.session_state:
            st.markdown(f'<div class="glass-card">{st.session_state.test_res}</div>', unsafe_allow_html=True)

# --- 3. 📝 ДААЛГАВАР ҮҮСГЭХ ---
elif menu == "📝 Даалгавар үүсгэх":
    st.title("📝 Даалгавар илгээх")
    with st.form("hw_form"):
        st.markdown('<div class="glass-card">', unsafe_allow_html=True)
        hw_grade = st.selectbox("Анги сонгох", [f"{i}-р анги" for i in range(1, 13)])
        hw_topic = st.text_input("Даалгаврын сэдэв")
        hw_desc = st.text_area("Зааварчилгаа")
        if st.form_submit_button("🚀 Даалгавар илгээх"):
            st.success(f"{hw_grade}-д даалгавар амжилттай илгээгдлээ.")
        st.markdown('</div>', unsafe_allow_html=True)

# --- 4. 🤖 AI ЧАТБОТ ---
elif menu == "🤖 AI Чатбот":
    st.title("🤖 AI Ухаалаг туслах")
    with st.expander("📎 Файл хавсаргах (PDF, Word, Excel, Зураг)"):
        up = st.file_uploader("Файл", type=['pdf', 'docx', 'xlsx', 'jpg', 'png'])
        f_context = ""
        if up:
            if up.name.endswith('.pdf'): f_context = extract_pdf_text(up)
            elif up.name.endswith('.docx'): f_context = extract_docx_text(up)
            elif up.name.endswith('.xlsx'):
                df = pd.read_excel(up); f_context = df.to_string(); st.dataframe(df.head())
            else: st.image(up, width=200); f_context = "[Зураг]"
            
    if "messages" not in st.session_state: st.session_state.messages = []
    for m in st.session_state.messages:
        with st.chat_message(m["role"]): st.markdown(m["content"])
    if p := st.chat_input("Асуултаа бичнэ үү..."):
        st.session_state.messages.append({"role": "user", "content": p})
        with st.chat_message("user"): st.markdown(p)
        res = requests.post("https://api.groq.com/openai/v1/chat/completions", 
                           headers={"Authorization": f"Bearer {st.secrets['GROQ_API_KEY']}"},
                           json={"model": "llama-3.3-70b-versatile", "messages": st.session_state.messages})
        ans = res.json()['choices'][0]['message']['content']
        with st.chat_message("assistant"): st.markdown(ans)
        st.session_state.messages.append({"role": "assistant", "content": ans})

# --- 5. 🌍 БОЛОВСРОЛ ---
elif menu == "🌍 Портал":
    st.title("🌍 Боловсролын Сайтууд")
    sites = {"🗺️ EduMap": "https://edumap.mn/", "🎥 Medle": "https://medle.mn/", "🎮 Eduten": "https://www.eduten.com/", "🇬🇧 Pearson": "https://englishconnect.pearson.com/", "📝 Unelgee": "https://unelgee.eec.mn/auth/login/", "📊 EEC": "https://www.eec.mn/", "Econtent": "https://econtent.edu.mn/"}
    t = st.tabs(list(sites.keys()))
    for i, (name, url) in enumerate(sites.items()):
        with t[i]: components.iframe(url, height=800)






